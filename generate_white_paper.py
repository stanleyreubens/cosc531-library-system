from docx import Document

# Create a new Word document
doc = Document()

# Title of the white paper
doc.add_heading("Library Management System White Paper", level=1)

# Introduction
doc.add_heading("Introduction", level=2)
doc.add_paragraph(
    "The Library Management System is a software solution designed to manage the borrowing and returning of books, "
    "while tracking student interactions dynamically. This project leverages Python, Flask for the frontend, and SQLite "
    "as the database backend. The system is intended to provide an efficient and user-friendly interface for library "
    "management operations."
)

# Problem Statement
doc.add_heading("Problem Statement", level=2)
doc.add_paragraph(
    "Libraries often face challenges in tracking book borrowings, managing returns, and ensuring the availability of resources. "
    "Students require easy access to library records to check their borrowed books and due dates. The goal of this project is to "
    "create a comprehensive library management system to address these challenges dynamically and with ease."
)

# System Design
doc.add_heading("System Design", level=2)

# Database Schema
doc.add_heading("Database Schema", level=3)
doc.add_paragraph(
    "The system uses a relational database (SQLite) to store information about books, students, and borrowing records. "
    "The database schema includes the following tables:"
)
doc.add_paragraph(
    "1. Books: Stores details of all available books, including title, author, and availability status."
)
doc.add_paragraph(
    "2. Students: Stores information about registered students."
)
doc.add_paragraph(
    "3. BorrowRecords: Tracks which books are borrowed, by whom, and the due dates for their return."
)

# Tools and Technologies
doc.add_heading("Tools and Technologies", level=3)
doc.add_paragraph(
    "1. Python: The core programming language used for backend development.\n"
    "2. Flask: A microframework for developing the web interface.\n"
    "3. SQLite: A lightweight database for storing application data.\n"
    "4. HTML/CSS: Used for the frontend UI.\n"
    "5. Pandas: For importing and processing book data from CSV files.\n"
)

# System Features
doc.add_heading("System Features", level=2)
doc.add_paragraph(
    "1. Dynamic Borrowing and Returning: Students can borrow or return books interactively via the web interface.\n"
    "2. Borrowed Book Lookup: Provides the ability to query which books a student has borrowed.\n"
    "3. User-Friendly Interface: A simple and intuitive web application for managing library operations.\n"
    "4. Duplicate Handling: Ensures that books cannot be borrowed multiple times until returned.\n"
    "5. Flash Notifications: Real-time feedback for user actions like borrowing or returning books.\n"
)

# System Workflow
doc.add_heading("System Workflow", level=2)
doc.add_paragraph(
    "1. Students can log in and search for available books.\n"
    "2. Books can be borrowed by specifying the student name and book title.\n"
    "3. Borrowed books can be returned, updating their availability in the database.\n"
    "4. The system dynamically tracks borrowing records, preventing duplicate borrowings."
)

# Conclusion
doc.add_heading("Conclusion", level=2)
doc.add_paragraph(
    "The Library Management System is a comprehensive solution for libraries seeking to digitize and streamline their operations. "
    "This project demonstrates the use of Flask and SQLite for building dynamic web applications with robust database integration."
)

# Save the Word document
file_path = "Library_Management_System_White_Paper.docx"
doc.save(file_path)

print(f"White paper saved to {file_path}")
